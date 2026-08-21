# ==================== app.py ====================
import os
import threading
import time
import schedule
from datetime import datetime, date, timedelta
from functools import wraps

from flask import (Flask, render_template, request, redirect, url_for,
                   jsonify, send_file, session, flash)
from flask_sqlalchemy import SQLAlchemy
from flask_session import Session
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

# ==================== CONFIGURATION ====================
app = Flask(__name__, template_folder='templates', static_folder='static')

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'shipping-company-secret-key-2024')

# ==================== SESSION CONFIGURATION ====================
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = '/tmp/flask_session'
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_KEY_PREFIX'] = 'adam_cargo_'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)

app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False

Session(app)

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

DATABASE_URL = os.environ.get('DATABASE_URL')
if not DATABASE_URL:
    DATABASE_URL = 'postgresql://adamscargo_postgress_user:NTnTZ1hYiCXJ1nrUnAyCsQym8Xm5ViUc@dpg-d9tjhrqd0e5s739brvl0-a/adamscargo_postgress'
    print("⚠️ تنبيه: لم يتم العثور على DATABASE_URL في البيئة، نستخدم الرابط المباشر")

if DATABASE_URL.startswith('postgres://'):
    DATABASE_URL = DATABASE_URL.replace('postgres://', 'postgresql://', 1)

app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 5,
    'pool_recycle': 3600,
    'pool_pre_ping': True,
    'pool_timeout': 30,
    'connect_args': {
        'connect_timeout': 10
    }
}

BACKUP_FOLDER_NAME = 'ShippingCompany_Backups'
BACKUP_FOLDER_ID = None

db = SQLAlchemy(app)

# ==================== MODELS ====================
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    full_name = db.Column(db.String(200))
    role = db.Column(db.String(20), default='user')
    active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.now)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Car(db.Model):
    __tablename__ = 'cars'
    id = db.Column(db.Integer, primary_key=True)
    plate_number = db.Column(db.String(20), unique=True, nullable=False)
    purchase_price = db.Column(db.Float, default=0)
    down_payment = db.Column(db.Float, default=0)
    bank_installment = db.Column(db.Float, default=0)
    remaining_bank = db.Column(db.Float, default=0)
    total_paid = db.Column(db.Float, default=0)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)

class Customer(db.Model):
    __tablename__ = 'customers'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False, unique=True)
    phone = db.Column(db.String(20))
    date_added = db.Column(db.DateTime, default=datetime.now)

class Trip(db.Model):
    __tablename__ = 'trips'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, nullable=False, default=date.today)
    car_id = db.Column(db.Integer, db.ForeignKey('cars.id'))
    driver_name = db.Column(db.String(100), nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customers.id'))
    from_location = db.Column(db.String(200))
    to_location = db.Column(db.String(200))
    nauloon = db.Column(db.Float, default=0)
    solar = db.Column(db.Float, default=0)
    expenses = db.Column(db.Float, default=0)
    driver_pay = db.Column(db.Float, default=0)
    net_profit = db.Column(db.Float, default=0)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))

    car = db.relationship('Car', backref='trips')
    customer = db.relationship('Customer', backref='trips')
    creator = db.relationship('User', backref='trips')

class Payment(db.Model):
    __tablename__ = 'payments'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, nullable=False, default=date.today)
    trip_id = db.Column(db.Integer, db.ForeignKey('trips.id'))
    customer_id = db.Column(db.Integer, db.ForeignKey('customers.id'))
    amount = db.Column(db.Float, default=0)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)

    trip = db.relationship('Trip', backref='payments')
    customer = db.relationship('Customer', backref='payments')

class Installment(db.Model):
    __tablename__ = 'installments'
    id = db.Column(db.Integer, primary_key=True)
    car_id = db.Column(db.Integer, db.ForeignKey('cars.id'))
    due_date = db.Column(db.Date)
    amount = db.Column(db.Float)
    notes = db.Column(db.Text, default='')
    paid = db.Column(db.Boolean, default=False)
    payment_date = db.Column(db.Date)
    created_at = db.Column(db.DateTime, default=datetime.now)

    car = db.relationship('Car', backref='installments')

class BankAccount(db.Model):
    __tablename__ = 'bank_accounts'
    id = db.Column(db.Integer, primary_key=True)
    bank_name = db.Column(db.String(100), nullable=False)
    account_number = db.Column(db.String(50))
    current_balance = db.Column(db.Float, default=0)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)

class BankTransaction(db.Model):
    __tablename__ = 'bank_transactions'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, nullable=False, default=date.today)
    type = db.Column(db.String(20), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    description = db.Column(db.Text)
    account_id = db.Column(db.Integer, db.ForeignKey('bank_accounts.id'), nullable=True)
    loan_id = db.Column(db.Integer, db.ForeignKey('bank_loans.id'), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.now)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))

    account = db.relationship('BankAccount', backref='transactions')
    loan = db.relationship('BankLoan', backref='transactions')
    creator = db.relationship('User', backref='bank_transactions')

class BankLoan(db.Model):
    __tablename__ = 'bank_loans'
    id = db.Column(db.Integer, primary_key=True)
    start_date = db.Column(db.Date, nullable=False, default=date.today)
    total_amount = db.Column(db.Float, nullable=False)
    monthly_installment = db.Column(db.Float, nullable=False)
    total_paid = db.Column(db.Float, default=0)
    remaining = db.Column(db.Float)
    description = db.Column(db.Text)
    account_id = db.Column(db.Integer, db.ForeignKey('bank_accounts.id'), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.now)

    account = db.relationship('BankAccount', backref='loans')

class LoanPayment(db.Model):
    __tablename__ = 'loan_payments'
    id = db.Column(db.Integer, primary_key=True)
    loan_id = db.Column(db.Integer, db.ForeignKey('bank_loans.id'))
    date = db.Column(db.Date, nullable=False, default=date.today)
    amount = db.Column(db.Float, nullable=False)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)

    loan = db.relationship('BankLoan', backref='payments')

class LandLoan(db.Model):
    __tablename__ = 'land_loans'
    id = db.Column(db.Integer, primary_key=True)
    land_price = db.Column(db.Float, default=0)
    total_loan_amount = db.Column(db.Float, default=0)
    total_paid = db.Column(db.Float, default=0)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)

class FinancialTransaction(db.Model):
    __tablename__ = 'financial_transactions'
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, nullable=False, default=date.today)
    person_name = db.Column(db.String(100), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    type = db.Column(db.String(10), nullable=False)
    description = db.Column(db.Text)
    bank_account_id = db.Column(db.Integer, db.ForeignKey('bank_accounts.id'), nullable=True)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_at = db.Column(db.DateTime, default=datetime.now)

    creator = db.relationship('User', backref='financial_transactions')
    bank_account = db.relationship('BankAccount', backref='financial_transactions')

# ==================== DECORATORS ====================
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            session['user_id'] = 1
            session['username'] = 'admin'
            session['full_name'] = 'مدير النظام'
            session['role'] = 'admin'
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            session['user_id'] = 1
            session['username'] = 'admin'
            session['full_name'] = 'مدير النظام'
            session['role'] = 'admin'
        if session.get('role') not in ['admin', 'root']:
            session['role'] = 'admin'
        return f(*args, **kwargs)
    return decorated_function

def root_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            session['user_id'] = 1
            session['username'] = 'admin'
            session['full_name'] = 'مدير النظام'
            session['role'] = 'root'
        if session.get('role') != 'root':
            session['role'] = 'root'
        return f(*args, **kwargs)
    return decorated_function

# ==================== GOOGLE DRIVE ====================
def get_drive():
    gauth = GoogleAuth()
    if os.path.exists('client_secrets.json'):
        gauth.LoadClientConfigFile('client_secrets.json')
    if os.path.exists('token.pickle'):
        gauth.LoadCredentialsFile('token.pickle')
    if gauth.credentials is None:
        gauth.LocalWebserverAuth()
        gauth.SaveCredentialsFile('token.pickle')
    elif gauth.access_token_expired:
        gauth.Refresh()
        gauth.SaveCredentialsFile('token.pickle')
    else:
        gauth.Authorize()
    return GoogleDrive(gauth)

def get_backup_folder(drive):
    global BACKUP_FOLDER_ID
    if BACKUP_FOLDER_ID:
        return BACKUP_FOLDER_ID
    lst = drive.ListFile({'q': f"title='{BACKUP_FOLDER_NAME}' and mimeType='application/vnd.google-apps.folder' and trashed=false"}).GetList()
    if lst:
        BACKUP_FOLDER_ID = lst[0]['id']
        return BACKUP_FOLDER_ID
    folder = drive.CreateFile({'title': BACKUP_FOLDER_NAME, 'mimeType': 'application/vnd.google-apps.folder'})
    folder.Upload()
    BACKUP_FOLDER_ID = folder['id']
    return BACKUP_FOLDER_ID

def backup_database():
    try:
        print("🔄 Backup...")
        ts = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        fn = f'backup_{ts}.xlsx'
        with pd.ExcelWriter(fn, engine='openpyxl') as w:
            trips = Trip.query.order_by(Trip.date.asc()).all()
            if trips:
                pd.DataFrame([{'التاريخ':str(t.date),'العربية':t.car.plate_number if t.car else '','السائق':t.driver_name,'العميل':t.customer.name if t.customer else '','من':t.from_location,'إلى':t.to_location,'النولون':t.nauloon,'السولار':t.solar,'المصاريف':t.expenses,'أجرة السائق':t.driver_pay,'الصافي':t.net_profit} for t in trips]).to_excel(w,sheet_name='الرحلات',index=False)
            cust = Customer.query.order_by(Customer.name.asc()).all()
            if cust:
                pd.DataFrame([{'الاسم':c.name,'التليفون':c.phone} for c in cust]).to_excel(w,sheet_name='العملاء',index=False)
            cars = Car.query.order_by(Car.plate_number.asc()).all()
            if cars:
                pd.DataFrame([{'اللوحة':c.plate_number,'سعر العربية':c.purchase_price,'المقدم':c.down_payment,'القسط':c.bank_installment,'المدفوع':c.total_paid,'المتبقي':c.remaining_bank} for c in cars]).to_excel(w,sheet_name='العربيات',index=False)
            pays = Payment.query.order_by(Payment.date.asc()).all()
            if pays:
                pd.DataFrame([{'التاريخ':str(p.date),'العميل':p.customer.name if p.customer else '','المبلغ':p.amount} for p in pays]).to_excel(w,sheet_name='الدفعات',index=False)
            insts = Installment.query.order_by(Installment.due_date.asc()).all()
            if insts:
                pd.DataFrame([{'العربية':i.car.plate_number if i.car else '','الاستحقاق':str(i.due_date),'المبلغ':i.amount,'تم':i.paid,'ملاحظات':i.notes} for i in insts]).to_excel(w,sheet_name='الأقساط',index=False)
            btx = BankTransaction.query.order_by(BankTransaction.date.asc()).all()
            if btx:
                pd.DataFrame([{'التاريخ':str(b.date),'النوع':b.type,'المبلغ':b.amount,'الوصف':b.description,'الحساب':b.account.bank_name if b.account else ''} for b in btx]).to_excel(w,sheet_name='البنك',index=False)
            ftx = FinancialTransaction.query.order_by(FinancialTransaction.date.asc()).all()
            if ftx:
                pd.DataFrame([{'التاريخ':str(f.date),'الشخص':f.person_name,'النوع':f.type,'المبلغ':f.amount,'الوصف':f.description,'الحساب':f.bank_account.bank_name if f.bank_account else ''} for f in ftx]).to_excel(w,sheet_name='معاملات_مالية',index=False)
        drive = get_drive()
        fid = get_backup_folder(drive)
        f = drive.CreateFile({'title':fn,'parents':[{'id':fid}],'mimeType':'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'})
        f.SetContentFile(fn)
        f.Upload()
        cutoff = datetime.now() - timedelta(days=30)
        for fl in drive.ListFile({'q':f"'{fid}' in parents and trashed=false"}).GetList():
            try:
                cd = datetime.strptime(fl['createdDate'],"%Y-%m-%dT%H:%M:%S.%fZ")
                if cd < cutoff: fl.Delete()
            except: pass
        os.remove(fn)
        print(f"✅ Backup: {fn}")
    except Exception as e:
        print(f"❌ Backup error: {e}")

def scheduler_loop():
    schedule.every().day.at("23:00").do(backup_database)
    schedule.every(6).hours.do(backup_database)
    while True:
        schedule.run_pending()
        time.sleep(60)

# ==================== أداة لجلب متغيرات الداشبورد ====================
def get_dashboard_context():
    tt = Trip.query.count()
    tcust = Customer.query.count()
    tcar = Car.query.count()
    today_tr = Trip.query.filter_by(date=date.today()).all()
    today_net = sum(t.net_profit for t in today_tr)
    today_nau = sum(t.nauloon for t in today_tr)
    bank_accounts = BankAccount.query.all()
    total_bank_balance = sum(a.current_balance for a in bank_accounts)
    loans = BankLoan.query.filter(BankLoan.remaining > 0).all()
    loan_rem = sum(l.remaining for l in loans)
    recent = Trip.query.order_by(Trip.date.desc()).limit(10).all()
    pending = []
    for c in Customer.query.all():
        tn = db.session.query(db.func.sum(Trip.nauloon)).filter(Trip.customer_id==c.id).scalar() or 0
        tp = db.session.query(db.func.sum(Payment.amount)).filter(Payment.customer_id==c.id).scalar() or 0
        r = tn - tp
        if r > 0: pending.append({'customer':c,'remaining':r})
    
    try:
        upcoming = Installment.query.filter(Installment.paid==False, Installment.due_date>=date.today()).order_by(Installment.due_date.asc()).limit(5).all()
    except:
        upcoming = []
    
    return {
        'total_trips': tt,
        'total_customers': tcust,
        'total_cars': tcar,
        'today_net': today_net,
        'today_nauloon': today_nau,
        'today_trips_count': len(today_tr),
        'total_bank_balance': total_bank_balance,
        'total_loan_remaining': loan_rem,
        'recent_trips': recent,
        'pending_payments': pending,
        'upcoming_installments': upcoming,
        'bank_accounts': bank_accounts
    }

# ==================== AUTH ====================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        u = request.form.get('username','').strip()
        p = request.form.get('password','')
        user = User.query.filter_by(username=u, active=True).first()
        if user and user.check_password(p):
            session.clear()
            session['user_id'] = user.id
            session['username'] = user.username
            session['full_name'] = user.full_name
            session['role'] = user.role
            session.permanent = True
            flash(f'مرحباً {user.full_name}!','success')
            return redirect('/dashboard')
        flash('خطأ في الدخول','danger')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('تم الخروج','info')
    return redirect(url_for('login'))

# ==================== DASHBOARD ====================
@app.route('/')
@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', **get_dashboard_context())

# ==================== ADMIN & USER MANAGEMENT ====================
@app.route('/users')
@admin_required
def users():
    if session.get('role') == 'root':
        all_users = User.query.order_by(User.full_name.asc()).all()
    else:
        all_users = User.query.filter(User.role != 'root').order_by(User.full_name.asc()).all()
    return render_template('users.html', users=all_users)

@app.route('/users/add', methods=['GET','POST'])
@admin_required
def add_user():
    if request.method == 'POST':
        un = request.form.get('username','').strip()
        if User.query.filter_by(username=un).first():
            flash('اسم المستخدم موجود بالفعل','danger')
            return redirect(url_for('add_user'))
        u = User(
            username=un,
            full_name=request.form.get('full_name','').strip(),
            role=request.form.get('role','user')
        )
        u.set_password(request.form.get('password',''))
        db.session.add(u)
        db.session.commit()
        flash('تمت الإضافة','success')
        return redirect(url_for('users'))
    return render_template('add_user.html')

@app.route('/api/users/<int:uid>/delete', methods=['POST'])
@admin_required
def delete_user(uid):
    u = User.query.get_or_404(uid)
    if u.role == 'admin' and User.query.filter_by(role='admin').count() <= 1:
        flash('لا يمكن حذف آخر أدمن','danger')
        return redirect(url_for('users'))
    db.session.delete(u)
    db.session.commit()
    flash('تم الحذف','success')
    return redirect(url_for('users'))

@app.route('/api/users/<int:uid>/reset-password', methods=['POST'])
@admin_required
def reset_user_password(uid):
    u = User.query.get_or_404(uid)
    u.set_password(request.form.get('new_password',''))
    db.session.commit()
    flash(f'تم تغيير كلمة المرور لـ {u.full_name}','success')
    return redirect(url_for('users'))

# ==================== LAND LOAN (ROOT ONLY) ====================
@app.route('/land')
@root_required
def land_loan():
    land = LandLoan.query.first()
    if not land:
        land = LandLoan(land_price=0, total_loan_amount=0, total_paid=0)
        db.session.add(land)
        db.session.commit()
    remaining = land.total_loan_amount - land.total_paid
    return render_template('land_loan.html', land=land, remaining=remaining)

@app.route('/api/land/update', methods=['POST'])
@root_required
def update_land():
    try:
        land = LandLoan.query.first()
        if not land:
            land = LandLoan()
            db.session.add(land)
        land.land_price = float(request.form.get('land_price', 0))
        land.total_loan_amount = float(request.form.get('total_loan_amount', 0))
        land.notes = request.form.get('notes', '')
        db.session.commit()
        flash('تم تحديث بيانات الأرض','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('land_loan'))

@app.route('/api/land/pay', methods=['POST'])
@root_required
def pay_land():
    try:
        land = LandLoan.query.first()
        if not land:
            flash('يجب إدخال بيانات الأرض أولاً','danger')
            return redirect(url_for('land_loan'))
        amt = float(request.form.get('amount', 0))
        remaining = land.total_loan_amount - land.total_paid
        if amt > remaining:
            flash(f'المبلغ أكبر من المتبقي! المتبقي: {remaining}','danger')
            return redirect(url_for('land_loan'))
        land.total_paid += amt
        db.session.commit()
        flash(f'تم دفع {amt} للأرض. المتبقي: {remaining - amt}','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('land_loan'))

# ==================== FINANCIAL TRANSACTIONS ====================
@app.route('/transactions')
@login_required
def financial_transactions():
    txns = FinancialTransaction.query.order_by(FinancialTransaction.date.desc()).all()
    total_given = db.session.query(db.func.sum(FinancialTransaction.amount)).filter(FinancialTransaction.type=='given').scalar() or 0
    total_received = db.session.query(db.func.sum(FinancialTransaction.amount)).filter(FinancialTransaction.type=='received').scalar() or 0
    bank_accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    return render_template('transactions.html', transactions=txns, total_given=total_given, total_received=total_received, bank_accounts=bank_accounts, date=date.today())

@app.route('/transactions/<person_name>')
@login_required
def person_transactions(person_name):
    txns = FinancialTransaction.query.filter_by(person_name=person_name).order_by(FinancialTransaction.date.desc()).all()
    total_given = db.session.query(db.func.sum(FinancialTransaction.amount)).filter(FinancialTransaction.person_name==person_name, FinancialTransaction.type=='given').scalar() or 0
    total_received = db.session.query(db.func.sum(FinancialTransaction.amount)).filter(FinancialTransaction.person_name==person_name, FinancialTransaction.type=='received').scalar() or 0
    return render_template('person_transactions.html', 
                           person_name=person_name, 
                           transactions=txns, 
                           total_given=total_given, 
                           total_received=total_received)

@app.route('/api/transactions/add', methods=['POST'])
@admin_required
def add_transaction():
    try:
        bank_account_id = request.form.get('bank_account_id') or None
        amount = float(request.form['amount'])
        txn_type = request.form['type']
        txn = FinancialTransaction(
            date=datetime.strptime(request.form['date'], '%Y-%m-%d').date(),
            person_name=request.form['person_name'],
            amount=amount,
            type=txn_type,
            description=request.form.get('description', ''),
            bank_account_id=bank_account_id if bank_account_id else None,
            created_by=session['user_id']
        )
        db.session.add(txn)
        db.session.flush()
        if bank_account_id:
            account = BankAccount.query.get(bank_account_id)
            if account:
                if txn_type == 'received':
                    account.current_balance += amount
                    db.session.add(BankTransaction(
                        date=txn.date,
                        type='deposit',
                        amount=amount,
                        description=f'استلام من {txn.person_name}',
                        account_id=bank_account_id,
                        created_by=session['user_id']
                    ))
                elif txn_type == 'given':
                    account.current_balance -= amount
                    db.session.add(BankTransaction(
                        date=txn.date,
                        type='withdraw',
                        amount=amount,
                        description=f'دفع إلى {txn.person_name}',
                        account_id=bank_account_id,
                        created_by=session['user_id']
                    ))
        db.session.commit()
        flash('تمت إضافة المعاملة المالية','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('financial_transactions'))

@app.route('/api/transactions/<int:tid>/edit', methods=['POST'])
@admin_required
def edit_transaction(tid):
    try:
        txn = FinancialTransaction.query.get_or_404(tid)
        txn.person_name = request.form.get('person_name', txn.person_name)
        txn.amount = float(request.form.get('amount', txn.amount) or 0)
        txn.type = request.form.get('type', txn.type)
        txn.description = request.form.get('description', txn.description)
        txn.date = datetime.strptime(request.form.get('date', str(txn.date)), '%Y-%m-%d').date()
        db.session.commit()
        flash('تم تعديل المعاملة بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('financial_transactions'))

@app.route('/api/transactions/<int:tid>/delete', methods=['POST'])
@admin_required
def delete_transaction(tid):
    try:
        txn = FinancialTransaction.query.get_or_404(tid)
        db.session.delete(txn)
        db.session.commit()
        flash('تم حذف المعاملة بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('financial_transactions'))

# ==================== TRIPS ====================
@app.route('/trips/add', methods=['GET','POST'])
@login_required
def add_trip():
    if request.method == 'POST':
        try:
            pn = request.form.get('plate_number','').strip()
            dn = request.form.get('driver_name','').strip()
            cn = request.form.get('customer_name','').strip()
            cp = request.form.get('customer_phone','').strip()
            car = Car.query.filter_by(plate_number=pn).first()
            if not car and pn:
                car = Car(plate_number=pn)
                db.session.add(car)
                db.session.flush()
            cust = Customer.query.filter_by(name=cn).first()
            if not cust and cn:
                cust = Customer(name=cn, phone=cp)
                db.session.add(cust)
                db.session.flush()
            nau = float(request.form.get('nauloon', 0) or 0)
            sol = float(request.form.get('solar', 0) or 0)
            exp = float(request.form.get('expenses', 0) or 0)
            dp = float(request.form.get('driver_pay', 0) or 0)
            net = nau - sol - exp - dp
            trip_date = datetime.strptime(request.form.get('date', str(date.today())), '%Y-%m-%d').date()
            trip = Trip(
                date=trip_date,
                car_id=car.id if car else None,
                driver_name=dn,
                customer_id=cust.id if cust else None,
                from_location=request.form.get('from_location',''),
                to_location=request.form.get('to_location',''),
                nauloon=nau,
                solar=sol,
                expenses=exp,
                driver_pay=dp,
                net_profit=net,
                notes=request.form.get('notes',''),
                created_by=session['user_id']
            )
            db.session.add(trip)
            db.session.flush()
            paid = float(request.form.get('paid_now', 0) or 0)
            if paid > 0:
                db.session.add(Payment(
                    date=trip_date,
                    trip_id=trip.id,
                    customer_id=cust.id if cust else None,
                    amount=paid,
                    notes='دفعة مع الرحلة'
                ))
            db.session.commit()
            flash('تمت الإضافة بنجاح', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ: {str(e)}', 'danger')
            print(f"Error: {e}")
        return redirect(url_for('add_trip'))
    return render_template(
        'add_trip.html',
        cars=Car.query.order_by(Car.plate_number.asc()).all(),
        customers=Customer.query.order_by(Customer.name.asc()).all(),
        today=date.today()
    )

@app.route('/trips')
@login_required
def trips_list():
    trips = Trip.query.order_by(Trip.date.asc()).all()
    return render_template('trips.html', trips=trips)

@app.route('/trips/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_trip(id):
    trip = Trip.query.get_or_404(id)
    if session.get('role') not in ['admin', 'root'] and trip.created_by != session['user_id']:
        flash('ليس لديك صلاحية لتعديل هذه الرحلة', 'danger')
        return redirect(url_for('trips_list'))
    if request.method == 'POST':
        try:
            trip.date = datetime.strptime(request.form['date'], '%Y-%m-%d').date()
            trip.driver_name = request.form['driver_name']
            trip.from_location = request.form.get('from_location', '')
            trip.to_location = request.form.get('to_location', '')
            trip.nauloon = float(request.form['nauloon'])
            trip.solar = float(request.form.get('solar', 0))
            trip.expenses = float(request.form.get('expenses', 0))
            trip.driver_pay = float(request.form.get('driver_pay', 0))
            trip.net_profit = trip.nauloon - trip.solar - trip.expenses - trip.driver_pay
            trip.notes = request.form.get('notes', '')
            db.session.commit()
            flash('تم تعديل الرحلة بنجاح', 'success')
            return redirect(url_for('trips_list'))
        except Exception as e:
            db.session.rollback()
            flash(f'حدث خطأ أثناء التعديل: {str(e)}', 'danger')
    cars = Car.query.order_by(Car.plate_number).all()
    customers = Customer.query.order_by(Customer.name).all()
    return render_template('edit_trip.html', trip=trip, cars=cars, customers=customers)

@app.route('/api/trips/<int:tid>/delete', methods=['POST'])
@admin_required
def delete_trip(tid):
    t = Trip.query.get_or_404(tid)
    Payment.query.filter_by(trip_id=tid).delete()
    db.session.delete(t)
    db.session.commit()
    flash('تم الحذف','success')
    return redirect(url_for('trips_list'))

# ==================== CUSTOMERS ====================
@app.route('/customers')
@login_required
def customers():
    sm = []
    for c in Customer.query.order_by(Customer.name.asc()).all():
        tn = db.session.query(db.func.sum(Trip.nauloon)).filter(Trip.customer_id==c.id).scalar() or 0
        tp = db.session.query(db.func.sum(Payment.amount)).filter(Payment.customer_id==c.id).scalar() or 0
        sm.append({'customer':c,'total_nauloon':tn,'total_paid':tp,'remaining':tn-tp})
    return render_template('customers.html', customers_summary=sm)

@app.route('/customers/<int:cid>')
@login_required
def customer_report(cid):
    c = Customer.query.get_or_404(cid)
    tr = Trip.query.filter_by(customer_id=cid).order_by(Trip.date.asc()).all()
    ps = Payment.query.filter_by(customer_id=cid).order_by(Payment.date.asc()).all()
    tn = sum(t.nauloon for t in tr)
    tp = sum(p.amount for p in ps)
    return render_template('customer_report.html', customer=c, trips=tr, payments=ps, total_nauloon=tn, total_paid=tp, remaining=tn-tp)

@app.route('/api/customers/<int:cid>/delete', methods=['POST'])
@admin_required
def delete_customer(cid):
    c = Customer.query.get_or_404(cid)
    Payment.query.filter_by(customer_id=cid).delete()
    Trip.query.filter_by(customer_id=cid).delete()
    db.session.delete(c)
    db.session.commit()
    flash('تم الحذف','success')
    return redirect(url_for('customers'))

@app.route('/payments/add', methods=['POST'])
@login_required
def add_payment():
    cid = request.form.get('customer_id')
    tid = request.form.get('trip_id')
    if not cid:
        flash('يرجى اختيار العميل','danger')
        return redirect(request.referrer or url_for('customers'))
    amt = float(request.form.get('amount', 0) or 0)
    db.session.add(Payment(
        date=datetime.strptime(request.form['date'],'%Y-%m-%d').date(),
        trip_id=tid if tid else None,
        customer_id=cid,
        amount=amt,
        notes=request.form.get('notes','')
    ))
    db.session.commit()
    flash('تمت الإضافة','success')
    return redirect(url_for('customer_report', cid=cid))

# ==================== CARS & INSTALLMENTS ====================
@app.route('/cars')
@login_required
def cars():
    cars_list = Car.query.order_by(Car.plate_number.asc()).all()
    bank_accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    return render_template('cars.html', cars=cars_list, bank_accounts=bank_accounts)

@app.route('/cars/<int:cid>')
@login_required
def car_report(cid):
    c = Car.query.get_or_404(cid)
    tr = Trip.query.filter_by(car_id=cid).order_by(Trip.date.asc()).all()
    drivers = {}
    for t in tr:
        drivers[t.driver_name] = drivers.get(t.driver_name, 0) + 1
    try:
        insts = Installment.query.filter_by(car_id=cid).order_by(Installment.due_date.asc()).all()
    except:
        insts = []
    bank_accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    return render_template('car_report.html', car=c, trips=tr, trip_count=len(tr),
                           total_nauloon=sum(t.nauloon for t in tr),
                           total_solar=sum(t.solar for t in tr),
                           total_expenses=sum(t.expenses for t in tr),
                           total_driver_pay=sum(t.driver_pay for t in tr),
                           total_net=sum(t.net_profit for t in tr),
                           drivers=drivers, installments=insts, bank_accounts=bank_accounts)

@app.route('/api/cars/add', methods=['POST'])
@admin_required
def add_car():
    try:
        plate = request.form.get('plate_number', '').strip()
        if Car.query.filter_by(plate_number=plate).first():
            flash('رقم اللوحة موجود بالفعل','danger')
            return redirect(url_for('cars'))
        
        purchase_price = float(request.form.get('purchase_price', 0) or 0)
        down_payment = float(request.form.get('down_payment', 0) or 0)
        remaining = purchase_price - down_payment
        
        car = Car(
            plate_number=plate,
            purchase_price=purchase_price,
            down_payment=down_payment,
            bank_installment=float(request.form.get('bank_installment', 0) or 0),
            remaining_bank=remaining,
            total_paid=down_payment,
            notes=request.form.get('notes', '')
        )
        db.session.add(car)
        db.session.commit()
        flash('تم إضافة العربية بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('cars'))

@app.route('/api/cars/<int:cid>/edit', methods=['POST'])
@admin_required
def edit_car(cid):
    try:
        car = Car.query.get_or_404(cid)
        car.plate_number = request.form.get('plate_number', car.plate_number).strip()
        car.purchase_price = float(request.form.get('purchase_price', car.purchase_price) or 0)
        car.down_payment = float(request.form.get('down_payment', car.down_payment) or 0)
        car.bank_installment = float(request.form.get('bank_installment', car.bank_installment) or 0)
        car.remaining_bank = float(request.form.get('remaining_bank', car.remaining_bank) or 0)
        car.total_paid = float(request.form.get('total_paid', car.total_paid) or 0)
        car.notes = request.form.get('notes', car.notes)
        db.session.commit()
        flash('تم تعديل العربية بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('cars'))

@app.route('/api/cars/<int:cid>/delete', methods=['POST'])
@admin_required
def delete_car(cid):
    try:
        c = Car.query.get_or_404(cid)
        Trip.query.filter_by(car_id=cid).delete()
        Installment.query.filter_by(car_id=cid).delete()
        db.session.delete(c)
        db.session.commit()
        flash('تم الحذف','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('cars'))

@app.route('/api/installments/add', methods=['POST'])
@admin_required
def add_installment():
    try:
        cid = request.form.get('car_id')
        if not cid:
            flash('يرجى اختيار العربية','danger')
            return redirect(url_for('cars'))
        
        car = Car.query.get(cid)
        if not car:
            flash('العربية غير موجودة','danger')
            return redirect(url_for('cars'))
        
        amt = float(request.form.get('amount', 0) or 0)
        account_id = request.form.get('account_id')
        
        inst = Installment(
            car_id=cid,
            due_date=datetime.strptime(request.form['due_date'],'%Y-%m-%d').date(),
            amount=amt,
            notes=request.form.get('notes', '')
        )
        
        db.session.add(inst)
        db.session.flush()
        
        if account_id:
            account = BankAccount.query.get(account_id)
            if account:
                account.current_balance -= amt
                db.session.add(BankTransaction(
                    date=inst.due_date,
                    type='withdraw',
                    amount=amt,
                    description=f'قسط عربية {car.plate_number}',
                    account_id=account_id,
                    created_by=session['user_id']
                ))
        
        db.session.commit()
        flash('تم إضافة القسط بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('cars'))

@app.route('/api/installments/<int:iid>/edit', methods=['POST'])
@admin_required
def edit_installment(iid):
    try:
        inst = Installment.query.get_or_404(iid)
        old_amount = inst.amount
        inst.amount = float(request.form.get('amount', inst.amount) or 0)
        inst.due_date = datetime.strptime(request.form.get('due_date', str(inst.due_date)), '%Y-%m-%d').date()
        inst.notes = request.form.get('notes', '')
        
        car = Car.query.get(inst.car_id)
        if car:
            car.remaining_bank = car.remaining_bank - old_amount + inst.amount
        
        db.session.commit()
        flash('تم تعديل القسط بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('car_report', cid=inst.car_id))

@app.route('/api/installments/<int:iid>/delete', methods=['POST'])
@admin_required
def delete_installment(iid):
    try:
        inst = Installment.query.get_or_404(iid)
        car_id = inst.car_id
        car = Car.query.get(car_id)
        if car and not inst.paid:
            car.remaining_bank -= inst.amount
        db.session.delete(inst)
        db.session.commit()
        flash('تم حذف القسط بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('car_report', cid=car_id))

@app.route('/api/installments/<int:iid>/pay', methods=['POST'])
@admin_required
def pay_installment(iid):
    try:
        inst = Installment.query.get_or_404(iid)
        inst.paid = True
        inst.payment_date = date.today()
        car = Car.query.get(inst.car_id)
        if car:
            car.remaining_bank -= inst.amount
            car.total_paid += inst.amount
            if car.remaining_bank <= 0:
                car.remaining_bank = 0
        account_id = request.form.get('account_id')
        if account_id:
            account = BankAccount.query.get(account_id)
            if account:
                account.current_balance -= inst.amount
                db.session.add(BankTransaction(
                    date=date.today(),
                    type='withdraw',
                    amount=inst.amount,
                    description=f'سداد قسط عربية {car.plate_number}',
                    account_id=account_id,
                    created_by=session['user_id']
                ))
        db.session.commit()
        flash('تم دفع القسط بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('car_report', cid=inst.car_id))

# ==================== BANK ACCOUNTS ====================
@app.route('/bank/accounts')
@admin_required
def bank_accounts():
    accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    total_balance = sum(a.current_balance for a in accounts)
    return render_template('bank_accounts.html', accounts=accounts, total_balance=total_balance)

@app.route('/api/bank/accounts/add', methods=['POST'])
@admin_required
def add_bank_account():
    try:
        bank_name = request.form.get('bank_name', '').strip()
        if not bank_name:
            flash('يرجى إدخال اسم البنك','danger')
            return redirect(url_for('bank_accounts'))
        account = BankAccount(
            bank_name=bank_name,
            account_number=request.form.get('account_number', '').strip(),
            current_balance=float(request.form.get('current_balance', 0) or 0),
            notes=request.form.get('notes', '')
        )
        db.session.add(account)
        db.session.commit()
        flash(f'تم إضافة حساب {bank_name} بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_accounts'))

@app.route('/api/bank/accounts/<int:aid>/edit', methods=['POST'])
@admin_required
def edit_bank_account(aid):
    try:
        account = BankAccount.query.get_or_404(aid)
        account.bank_name = request.form.get('bank_name', account.bank_name).strip()
        account.account_number = request.form.get('account_number', account.account_number).strip()
        account.current_balance = float(request.form.get('current_balance', account.current_balance) or 0)
        account.notes = request.form.get('notes', account.notes)
        db.session.commit()
        flash('تم تحديث الحساب','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_accounts'))

@app.route('/api/bank/accounts/<int:aid>/deposit', methods=['POST'])
@admin_required
def deposit_bank_account(aid):
    try:
        account = BankAccount.query.get_or_404(aid)
        amount = float(request.form.get('amount', 0) or 0)
        account.current_balance += amount
        db.session.add(BankTransaction(
            date=datetime.strptime(request.form.get('date', str(date.today())), '%Y-%m-%d').date(),
            type='deposit',
            amount=amount,
            description=request.form.get('description', 'إيداع مبلغ'),
            account_id=aid,
            created_by=session['user_id']
        ))
        db.session.commit()
        flash(f'تم إيداع {amount} في {account.bank_name}','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_accounts'))

@app.route('/api/bank/accounts/<int:aid>/delete', methods=['POST'])
@admin_required
def delete_bank_account(aid):
    try:
        account = BankAccount.query.get_or_404(aid)
        if account.transactions:
            flash('لا يمكن حذف الحساب لوجود معاملات مرتبطة به','danger')
            return redirect(url_for('bank_accounts'))
        db.session.delete(account)
        db.session.commit()
        flash('تم حذف الحساب','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_accounts'))

# ==================== BANK ====================
@app.route('/bank')
@admin_required
def bank():
    accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    total_balance = sum(a.current_balance for a in accounts)
    loans = BankLoan.query.filter(BankLoan.remaining > 0).all()
    loan_rem = sum(l.remaining for l in loans)
    recent_transactions = BankTransaction.query.order_by(BankTransaction.date.asc()).all()
    return render_template('bank.html',
                           accounts=accounts,
                           total_balance=total_balance,
                           total_loan_remaining=loan_rem,
                           active_loans=loans,
                           recent_transactions=recent_transactions,
                           date=date.today())

@app.route('/bank/transaction/add', methods=['POST'])
@admin_required
def add_bank_transaction():
    try:
        account_id = request.form.get('account_id')
        if not account_id:
            flash('يرجى اختيار الحساب البنكي','danger')
            return redirect(url_for('bank'))
        account = BankAccount.query.get(account_id)
        if not account:
            flash('الحساب غير موجود','danger')
            return redirect(url_for('bank'))
        txn_type = request.form['type']
        amount = float(request.form['amount'] or 0)
        if txn_type == 'deposit':
            account.current_balance += amount
        elif txn_type == 'withdraw':
            account.current_balance -= amount
        db.session.add(BankTransaction(
            date=datetime.strptime(request.form['date'],'%Y-%m-%d').date(),
            type=txn_type,
            amount=amount,
            description=request.form.get('description',''),
            account_id=account_id,
            created_by=session['user_id']
        ))
        db.session.commit()
        flash('تمت الإضافة','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank'))

@app.route('/api/bank/transactions/<int:tid>/delete', methods=['POST'])
@admin_required
def delete_bank_transaction(tid):
    try:
        txn = BankTransaction.query.get_or_404(tid)
        if txn.account:
            if txn.type == 'deposit':
                txn.account.current_balance -= txn.amount
            elif txn.type == 'withdraw':
                txn.account.current_balance += txn.amount
        db.session.delete(txn)
        db.session.commit()
        flash('تم حذف المعاملة البنكية بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank'))

@app.route('/bank/loans')
@admin_required
def bank_loans():
    loans = BankLoan.query.order_by(BankLoan.start_date.asc()).all()
    accounts = BankAccount.query.order_by(BankAccount.bank_name.asc()).all()
    return render_template('bank_loans.html', loans=loans, accounts=accounts, date=date)

@app.route('/bank/loans/add', methods=['POST'])
@admin_required
def add_bank_loan():
    try:
        amt = float(request.form['total_amount'] or 0)
        mon = float(request.form['monthly_installment'] or 0)
        account_id = request.form.get('account_id')
        if not account_id:
            flash('يرجى اختيار الحساب البنكي','danger')
            return redirect(url_for('bank_loans'))
        loan = BankLoan(
            start_date=datetime.strptime(request.form['start_date'],'%Y-%m-%d').date(),
            total_amount=amt,
            monthly_installment=mon,
            remaining=amt,
            description=request.form.get('description',''),
            account_id=account_id
        )
        db.session.add(loan)
        db.session.flush()
        account = BankAccount.query.get(account_id)
        if account:
            account.current_balance += amt
            db.session.add(BankTransaction(
                date=loan.start_date,
                type='deposit',
                amount=amt,
                description=f'قرض جديد - {loan.description or ""}',
                account_id=account_id,
                loan_id=loan.id,
                created_by=session['user_id']
            ))
        db.session.commit()
        flash('تمت إضافة القرض','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_loans'))

@app.route('/bank/loans/<int:lid>/edit', methods=['POST'])
@admin_required
def edit_bank_loan(lid):
    try:
        loan = BankLoan.query.get_or_404(lid)
        loan.total_amount = float(request.form.get('total_amount', loan.total_amount) or 0)
        loan.monthly_installment = float(request.form.get('monthly_installment', loan.monthly_installment) or 0)
        loan.description = request.form.get('description', '')
        account_id = request.form.get('account_id')
        if account_id:
            loan.account_id = account_id
        loan.remaining = loan.total_amount - loan.total_paid
        db.session.commit()
        flash('تم تعديل القرض بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_loans'))

@app.route('/bank/loans/<int:lid>/delete', methods=['POST'])
@admin_required
def delete_bank_loan(lid):
    try:
        loan = BankLoan.query.get_or_404(lid)
        LoanPayment.query.filter_by(loan_id=lid).delete()
        BankTransaction.query.filter_by(loan_id=lid).delete()
        db.session.delete(loan)
        db.session.commit()
        flash('تم حذف القرض بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_loans'))

@app.route('/bank/loans/<int:lid>/pay', methods=['POST'])
@admin_required
def pay_loan_installment(lid):
    try:
        loan = BankLoan.query.get_or_404(lid)
        amt = float(request.form['amount'] or 0)
        account_id = request.form.get('account_id')
        if not account_id:
            flash('يرجى اختيار الحساب البنكي','danger')
            return redirect(url_for('bank_loans'))
        db.session.add(LoanPayment(
            loan_id=lid,
            date=datetime.strptime(request.form['date'],'%Y-%m-%d').date(),
            amount=amt,
            notes=request.form.get('notes','')
        ))
        loan.total_paid += amt
        loan.remaining = loan.total_amount - loan.total_paid
        account = BankAccount.query.get(account_id)
        if account:
            account.current_balance -= amt
            db.session.add(BankTransaction(
                date=datetime.strptime(request.form['date'],'%Y-%m-%d').date(),
                type='withdraw',
                amount=amt,
                description=f'سداد قرض #{lid}',
                account_id=account_id,
                loan_id=lid,
                created_by=session['user_id']
            ))
        db.session.commit()
        flash('تم السداد','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('bank_loans'))

# ==================== REPORTS ====================
@app.route('/api/installments/<int:iid>/pay', methods=['POST'])
@admin_required
def pay_installment(iid):
    try:
        inst = Installment.query.get_or_404(iid)
        inst.paid = True
        
        # تعديل: استخدمي التاريخ المدخل من المستخدم
        payment_date_str = request.form.get('payment_date', '')
        if payment_date_str:
            inst.payment_date = datetime.strptime(payment_date_str, '%Y-%m-%d').date()
        else:
            inst.payment_date = date.today()
        
        car = Car.query.get(inst.car_id)
        if car:
            car.remaining_bank -= inst.amount
            car.total_paid += inst.amount
            if car.remaining_bank <= 0:
                car.remaining_bank = 0
        account_id = request.form.get('account_id')
        if account_id:
            account = BankAccount.query.get(account_id)
            if account:
                account.current_balance -= inst.amount
                db.session.add(BankTransaction(
                    date=inst.payment_date,
                    type='withdraw',
                    amount=inst.amount,
                    description=f'سداد قسط عربية {car.plate_number}',
                    account_id=account_id,
                    created_by=session['user_id']
                ))
        db.session.commit()
        flash('تم دفع القسط بنجاح','success')
    except Exception as e:
        db.session.rollback()
        flash(f'حدث خطأ: {str(e)}','danger')
    return redirect(url_for('car_report', cid=inst.car_id))
@app.route('/reports/custom')
@login_required
def custom_report():
    start_date_str = request.args.get('start_date', '')
    end_date_str = request.args.get('end_date', '')
    
    trips = []
    total_nauloon = 0
    total_solar = 0
    total_expenses = 0
    total_driver_pay = 0
    total_net = 0
    
    if start_date_str and end_date_str:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        trips = Trip.query.filter(Trip.date >= start_date, Trip.date <= end_date).order_by(Trip.date.asc()).all()
        total_nauloon = sum(t.nauloon for t in trips)
        total_solar = sum(t.solar for t in trips)
        total_expenses = sum(t.expenses for t in trips)
        total_driver_pay = sum(t.driver_pay for t in trips)
        total_net = sum(t.net_profit for t in trips)
    
    return render_template('custom_report.html',
                           trips=trips,
                           start_date=start_date_str,
                           end_date=end_date_str,
                           total_nauloon=total_nauloon,
                           total_solar=total_solar,
                           total_expenses=total_expenses,
                           total_driver_pay=total_driver_pay,
                           total_net=total_net)

@app.route('/reports/custom/export')
@login_required
def export_custom_report():
    start_date_str = request.args.get('start_date', '')
    end_date_str = request.args.get('end_date', '')
    
    if not start_date_str or not end_date_str:
        flash('يرجى تحديد التاريخين','danger')
        return redirect(url_for('custom_report'))
    
    start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    trips = Trip.query.filter(Trip.date >= start_date, Trip.date <= end_date).order_by(Trip.date.asc()).all()
    
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    title = doc.add_heading('ADAM CARGO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company = doc.add_heading('شركة آدم للشحن والنقل', 1)
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title = doc.add_heading(f'تقرير الرحلات من {start_date_str} إلى {end_date_str}', 2)
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_number = datetime.now().strftime('%Y%m%d%H%M%S')
    report_no = doc.add_paragraph(f'رقم التقرير: {report_number}')
    report_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_date = doc.add_paragraph(f'تاريخ الإنشاء: {date.today()}')
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_heading('ملخص التقرير', 3)
    doc.add_paragraph(f'عدد الرحلات: {len(trips)}')
    doc.add_paragraph(f'إجمالي النولون: {sum(t.nauloon for t in trips)}')
    doc.add_paragraph(f'إجمالي السولار: {sum(t.solar for t in trips)}')
    doc.add_paragraph(f'إجمالي المصاريف: {sum(t.expenses for t in trips)}')
    doc.add_paragraph(f'إجمالي أجرة السائق: {sum(t.driver_pay for t in trips)}')
    doc.add_paragraph(f'إجمالي الصافي: {sum(t.net_profit for t in trips)}')
    doc.add_paragraph('')
    doc.add_heading('تفاصيل الرحلات', 3)
    table = doc.add_table(rows=1, cols=11)
    table.style = 'Light Shading Accent 1'
    headers = ['#', 'التاريخ', 'العربية', 'السائق', 'العميل', 'من', 'إلى', 'النولون', 'السولار', 'المصاريف', 'الصافي']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for idx, trip in enumerate(trips, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(trip.date)
        row[2].text = trip.car.plate_number if trip.car else '-'
        row[3].text = trip.driver_name
        row[4].text = trip.customer.name if trip.customer else '-'
        row[5].text = trip.from_location or '-'
        row[6].text = trip.to_location or '-'
        row[7].text = str(trip.nauloon)
        row[8].text = str(trip.solar)
        row[9].text = str(trip.expenses)
        row[10].text = str(trip.net_profit)
    
    fn = f'trips_report_{start_date_str}_to_{end_date_str}.docx'
    doc.save(fn)
    return send_file(fn, as_attachment=True)

@app.route('/reports/daily')
@login_required
def daily_report():
    rd = request.args.get('date', date.today().isoformat())
    rd = datetime.strptime(rd, '%Y-%m-%d').date()
    tr = Trip.query.filter_by(date=rd).order_by(Trip.id.asc()).all()
    return render_template('daily_report.html', report_date=rd, trips=tr,
                           total_nauloon=sum(t.nauloon for t in tr),
                           total_solar=sum(t.solar for t in tr),
                           total_expenses=sum(t.expenses for t in tr),
                           total_driver_pay=sum(t.driver_pay for t in tr),
                           total_net=sum(t.net_profit for t in tr))

@app.route('/reports/monthly')
@login_required
def monthly_report():
    month = request.args.get('month', date.today().strftime('%Y-%m'))
    year, mon = map(int, month.split('-'))
    start_date = date(year, mon, 1)
    if mon == 12:
        end_date = date(year + 1, 1, 1)
    else:
        end_date = date(year, mon + 1, 1)
    tr = Trip.query.filter(Trip.date >= start_date, Trip.date < end_date).order_by(Trip.date.asc()).all()
    total_nauloon = sum(t.nauloon for t in tr)
    total_solar = sum(t.solar for t in tr)
    total_expenses = sum(t.expenses for t in tr)
    total_driver_pay = sum(t.driver_pay for t in tr)
    total_net = sum(t.net_profit for t in tr)
    daily_summary = {}
    for t in tr:
        d = t.date
        if d not in daily_summary:
            daily_summary[d] = {'trip_count': 0, 'nauloon': 0, 'solar': 0, 'expenses': 0, 'driver_pay': 0, 'net': 0}
        daily_summary[d]['trip_count'] += 1
        daily_summary[d]['nauloon'] += t.nauloon
        daily_summary[d]['solar'] += t.solar
        daily_summary[d]['expenses'] += t.expenses
        daily_summary[d]['driver_pay'] += t.driver_pay
        daily_summary[d]['net'] += t.net_profit
    sorted_daily = dict(sorted(daily_summary.items()))
    return render_template('monthly_report.html',
                           report_month=month, trips=tr,
                           total_nauloon=total_nauloon,
                           total_solar=total_solar,
                           total_expenses=total_expenses,
                           total_driver_pay=total_driver_pay,
                           total_net=total_net,
                           daily_summary=sorted_daily)

@app.route('/reports/daily/export/<rd>')
@login_required
def export_daily_report(rd):
    rd = datetime.strptime(rd, '%Y-%m-%d').date()
    tr = Trip.query.filter_by(date=rd).order_by(Trip.id.asc()).all()
    df = pd.DataFrame([{'التاريخ':str(t.date),'العربية':t.car.plate_number if t.car else '','السائق':t.driver_name,
                        'العميل':t.customer.name if t.customer else '','من':t.from_location,'إلى':t.to_location,
                        'النولون':t.nauloon,'السولار':t.solar,'المصاريف':t.expenses,
                        'أجرة السائق':t.driver_pay,'الصافي':t.net_profit} for t in tr])
    if not df.empty:
        tot = {'التاريخ':'الإجمالي','العربية':'','السائق':'','العميل':'','من':'','إلى':''}
        for col in ['النولون','السولار','المصاريف','أجرة السائق','الصافي']:
            tot[col] = df[col].sum()
        df.loc['الإجمالي'] = tot
    fn = f'trips_report_{rd}.xlsx'
    df.to_excel(fn, index=False, engine='openpyxl')
    return send_file(fn, as_attachment=True)

@app.route('/reports/monthly/export/<month>')
@login_required
def export_monthly_report(month):
    year, mon = map(int, month.split('-'))
    start_date = date(year, mon, 1)
    if mon == 12:
        end_date = date(year + 1, 1, 1)
    else:
        end_date = date(year, mon + 1, 1)
    tr = Trip.query.filter(Trip.date >= start_date, Trip.date < end_date).order_by(Trip.date.asc()).all()
    df = pd.DataFrame([{'التاريخ':str(t.date),'العربية':t.car.plate_number if t.car else '','السائق':t.driver_name,
                        'العميل':t.customer.name if t.customer else '','من':t.from_location,'إلى':t.to_location,
                        'النولون':t.nauloon,'السولار':t.solar,'المصاريف':t.expenses,
                        'أجرة السائق':t.driver_pay,'الصافي':t.net_profit} for t in tr])
    if not df.empty:
        tot = {'التاريخ':'الإجمالي','العربية':'','السائق':'','العميل':'','من':'','إلى':''}
        for col in ['النولون','السولار','المصاريف','أجرة السائق','الصافي']:
            tot[col] = df[col].sum()
        df.loc['الإجمالي'] = tot
    fn = f'trips_monthly_report_{month}.xlsx'
    df.to_excel(fn, index=False, engine='openpyxl')
    return send_file(fn, as_attachment=True)

@app.route('/api/backup', methods=['POST'])
@admin_required
def manual_backup():
    try:
        backup_database()
        flash('تم النسخ الاحتياطي','success')
    except Exception as e:
        flash(f'فشل: {e}','danger')
    return redirect(url_for('dashboard'))

# ==================== API ====================
@app.route('/api/cars')
@login_required
def api_cars():
    q = request.args.get('search','')
    cars = Car.query.filter(Car.plate_number.contains(q)).order_by(Car.plate_number.asc()).all()
    return jsonify([{'id':c.id,'plate_number':c.plate_number} for c in cars])

@app.route('/api/customers')
@login_required
def api_customers():
    q = request.args.get('search','')
    customers = Customer.query.filter(Customer.name.contains(q)).order_by(Customer.name.asc()).all()
    return jsonify([{'id':c.id,'name':c.name,'phone':c.phone or ''} for c in customers])

@app.route('/api/customer_trips')
@login_required
def api_customer_trips():
    cid = request.args.get('customer_id','')
    if cid:
        trips = Trip.query.filter_by(customer_id=cid).order_by(Trip.date.asc()).all()
        return jsonify([{'id':t.id,'date':str(t.date),'nauloon':t.nauloon,'from_location':t.from_location,'to_location':t.to_location}
                        for t in trips])
    return jsonify([])

# ==================== ERROR HANDLERS ====================
@app.errorhandler(404)
def not_found(e):
    return render_template('login.html'), 404

@app.errorhandler(500)
def internal_error(e):
    db.session.rollback()
    return render_template('login.html'), 500

# ==================== HEALTH CHECK ====================
@app.route('/health')
def health_check():
    return 'OK', 200

# ==================== INIT ====================
def init_db():
    with app.app_context():
        try:
            db.session.execute(db.text('ALTER TABLE cars ADD COLUMN IF NOT EXISTS purchase_price FLOAT DEFAULT 0'))
            db.session.execute(db.text('ALTER TABLE cars ADD COLUMN IF NOT EXISTS down_payment FLOAT DEFAULT 0'))
            db.session.execute(db.text('ALTER TABLE cars ADD COLUMN IF NOT EXISTS total_paid FLOAT DEFAULT 0'))
            db.session.commit()
            print("✅ تم إضافة الأعمدة الجديدة لجدول cars")
        except Exception as e:
            print(f"⚠️ خطأ في تحديث جدول cars: {e}")
            db.session.rollback()
        
        try:
            db.session.execute(db.text('ALTER TABLE installments ADD COLUMN IF NOT EXISTS notes TEXT DEFAULT \'\''))
            db.session.execute(db.text('ALTER TABLE installments ADD COLUMN IF NOT EXISTS paid BOOLEAN DEFAULT false'))
            db.session.execute(db.text('ALTER TABLE installments ADD COLUMN IF NOT EXISTS payment_date DATE'))
            db.session.execute(db.text('ALTER TABLE installments ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP'))
            db.session.commit()
            print("✅ تم تحديث جدول installments بنجاح")
        except Exception as e:
            print(f"⚠️ خطأ في تحديث جدول installments: {e}")
            db.session.rollback()
        
        db.create_all()

        if not User.query.filter_by(username='admin').first():
            a = User(username='admin', full_name='مدير النظام', role='admin')
            a.set_password('admin123')
            db.session.add(a)
            db.session.commit()
            print("✅ Admin: admin / admin123")

        if not User.query.filter_by(username='hany').first():
            r = User(username='hany', full_name='Hany (Owner)', role='root')
            r.set_password('Hany@2024Secure')
            db.session.add(r)
            db.session.commit()
            print("✅ Root user (hany) created")

        if BankAccount.query.count() == 0:
            db.session.add(BankAccount(bank_name='بنك مصر', account_number='', current_balance=0))
            db.session.add(BankAccount(bank_name='البنك الأهلي', account_number='', current_balance=0))
            db.session.commit()
            print("✅ تمت إضافة حسابين بنكيين افتراضيين")


# ==================== تشغيل مباشر ====================
init_db()
threading.Thread(target=scheduler_loop, daemon=True).start()
print("✅ Scheduler running")

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
